"""文档解析与分块。

职责：
1. 解析常见文档格式为纯文本
2. 文本清洗
3. 自适应语义分块（按段落/句子边界递归切分，控制 chunk 长度与重叠）
4. 基于内容哈希的重复片段去重
"""
from __future__ import annotations

import csv
import json
import mimetypes
import os
import re
import zipfile
from dataclasses import dataclass, field
from html.parser import HTMLParser
from typing import List
from xml.etree import ElementTree as ET

from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from utils.common import clean_text, get_logger, text_hash
from utils.config import settings

logger = get_logger("rag.loader")

SUPPORTED_TYPES = {
    "text": [".txt", ".md", ".csv", ".json", ".jsonl", ".xml", ".html", ".htm"],
    "pdf": [".pdf"],
    "office": [".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".rtf", ".odt", ".ods", ".odp"],
    "ebook": [".epub"],
    "image_ocr": [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif"],
}
UNSUPPORTED_TYPES = {}


def supported_types() -> dict:
    """返回当前解析器支持/暂不支持的格式清单。"""
    return {
        "supported": SUPPORTED_TYPES,
        "unsupported": UNSUPPORTED_TYPES,
        "ocr": {
            "enabled": settings.enable_ocr,
            "provider": settings.ocr_provider,
            "model": settings.ocr_model,
            "max_pdf_pages": settings.ocr_max_pages,
            "dpi": settings.ocr_dpi,
        },
        "notes": {
            "pdf": "优先提取 PDF 内置文本；若文本过少且开启 OCR，则渲染页面走视觉 OCR。",
            "legacy_office": ".doc/.xls/.ppt 依赖 Windows + 已安装 Microsoft Office。",
            "ocr": "图片与扫描 PDF OCR 默认关闭，需配置 ENABLE_OCR=true 与 OCR provider。",
        },
    }


@dataclass
class Chunk:
    """单个文本片段及其元数据。"""

    text: str
    source: str  # 来源文件名
    chunk_index: int
    content_hash: str = ""

    def __post_init__(self) -> None:
        if not self.content_hash:
            self.content_hash = text_hash(self.text)

    def metadata(self) -> dict:
        return {
            "source": self.source,
            "chunk_index": self.chunk_index,
            "content_hash": self.content_hash,
        }


@dataclass
class LoadResult:
    """一次加载的结果汇总。"""

    chunks: List[Chunk] = field(default_factory=list)
    skipped_duplicates: int = 0


# 中英文混排友好的分隔符优先级：段落 -> 换行 -> 中文句号/问号/感叹号 -> 英文标点 -> 空格
_SEPARATORS = ["\n\n", "\n", "。", "！", "？", "；", ". ", "! ", "? ", "; ", " ", ""]


def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # 个别页解析失败不应中断整篇
            logger.warning("PDF 某页解析失败 %s: %s", path, exc)
    text = "\n".join(pages)
    if len(clean_text(text)) < settings.ocr_pdf_min_text_chars and settings.enable_ocr:
        logger.info("PDF 文本层内容较少，尝试 OCR：%s", path)
        return _read_pdf_ocr(path)
    return text


def _read_image_ocr(path: str) -> str:
    if not settings.enable_ocr:
        raise ValueError("图片 OCR 未启用。请设置 ENABLE_OCR=true 并配置 OCR_PROVIDER/OCR_API_KEY。")
    from rag.ocr import get_ocr_provider

    mime_type = mimetypes.guess_type(path)[0] or "image/png"
    with open(path, "rb") as f:
        return get_ocr_provider().extract_image_bytes(f.read(), mime_type=mime_type)


def _read_pdf_ocr(path: str) -> str:
    if not settings.enable_ocr:
        raise ValueError("扫描版 PDF OCR 未启用。请设置 ENABLE_OCR=true 并配置 OCR_PROVIDER/OCR_API_KEY。")
    try:
        import fitz
    except Exception as exc:
        raise ValueError("扫描版 PDF OCR 需要安装 PyMuPDF：pip install PyMuPDF==1.24.14") from exc

    from rag.ocr import get_ocr_provider

    provider = get_ocr_provider()
    doc = fitz.open(path)
    total_pages = min(len(doc), settings.ocr_max_pages)
    zoom = settings.ocr_dpi / 72
    matrix = fitz.Matrix(zoom, zoom)
    parts: List[str] = []
    try:
        for page_index in range(total_pages):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            image = pix.tobytes("png")
            text = provider.extract_image_bytes(image, mime_type="image/png")
            if text:
                parts.append(f"[第 {page_index + 1} 页]\n{text}")
        if len(doc) > total_pages:
            logger.warning("PDF OCR 仅处理前 %d 页：%s", total_pages, path)
    finally:
        doc.close()
    return "\n\n".join(parts)


def _read_docx(path: str) -> str:
    """解析 Word .docx，提取段落与表格文本。"""
    from docx import Document

    doc = Document(path)
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text)
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_doc(path: str) -> str:
    """解析老版 Word .doc（二进制）。

    依赖本机安装的 Microsoft Word（通过 COM 自动化读取）。
    若未安装 Word 或调用失败，则提示用户另存为 .docx。
    """
    try:
        import pythoncom  # noqa: F401
        import win32com.client as win32
    except Exception as exc:
        raise ValueError(
            ".doc 解析需要 Windows + 已安装 Microsoft Word。"
            "建议用 Word 另存为 .docx 后再上传。"
        ) from exc

    import os

    import pythoncom

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
        text = doc.Content.Text
        doc.Close(False)
        return text or ""
    except Exception as exc:
        raise ValueError(
            f".doc 解析失败（请确认已安装 Microsoft Word）：{exc}。"
            "建议用 Word 另存为 .docx 后再上传。"
        ) from exc
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


class _HTMLTextExtractor(HTMLParser):
    """简单 HTML 文本提取器。"""

    def __init__(self) -> None:
        super().__init__()
        self.parts: List[str] = []

    def handle_data(self, data: str) -> None:
        if data and data.strip():
            self.parts.append(data.strip())

    def text(self) -> str:
        return "\n".join(self.parts)


def _read_html(path: str) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(_read_txt(path))
    return parser.text()


def _read_csv(path: str) -> str:
    raw = _read_txt(path)
    sample = raw[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel
    rows = []
    for row in csv.reader(raw.splitlines(), dialect):
        cells = [c.strip() for c in row if c and c.strip()]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _read_json(path: str) -> str:
    raw = _read_txt(path)
    try:
        data = json.loads(raw)
        return json.dumps(data, ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        # 兼容 JSONL：逐行解析，失败则保留原文。
        lines = []
        for line in raw.splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                lines.append(json.dumps(json.loads(line), ensure_ascii=False))
            except json.JSONDecodeError:
                return raw
        return "\n".join(lines)


def _read_xml(path: str) -> str:
    raw = _read_txt(path)
    try:
        root = ET.fromstring(raw)
        return "\n".join(t.strip() for t in root.itertext() if t and t.strip())
    except ET.ParseError:
        return re.sub(r"<[^>]+>", " ", raw)


def _read_rtf(path: str) -> str:
    text = _read_txt(path)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+\d* ?", " ", text)
    text = re.sub(r"[{}]", " ", text)
    return text


def _zip_xml_text(path: str, members: List[str] | None = None, suffixes: tuple[str, ...] = ()) -> str:
    """从 zip 容器中的 XML/HTML 文件提取文本。"""
    parts: List[str] = []
    with zipfile.ZipFile(path) as zf:
        names = members or [
            name for name in zf.namelist() if not suffixes or name.lower().endswith(suffixes)
        ]
        for name in names:
            try:
                raw = zf.read(name)
            except KeyError:
                continue
            try:
                root = ET.fromstring(raw)
                parts.extend(t.strip() for t in root.itertext() if t and t.strip())
            except ET.ParseError:
                try:
                    parser = _HTMLTextExtractor()
                    parser.feed(raw.decode("utf-8", errors="ignore"))
                    html_text = parser.text()
                    if html_text:
                        parts.append(html_text)
                except Exception:
                    continue
    return "\n".join(parts)


def _read_xlsx(path: str) -> str:
    """解析 .xlsx，提取共享字符串与工作表文本。"""
    shared: List[str] = []
    parts: List[str] = []
    with zipfile.ZipFile(path) as zf:
        if "xl/sharedStrings.xml" in zf.namelist():
            root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
            shared = ["".join(si.itertext()).strip() for si in root if "".join(si.itertext()).strip()]
        sheet_names = sorted(
            name for name in zf.namelist() if name.startswith("xl/worksheets/") and name.endswith(".xml")
        )
        for sheet_name in sheet_names:
            root = ET.fromstring(zf.read(sheet_name))
            rows = []
            for row in root.iter():
                if not row.tag.endswith("row"):
                    continue
                cells = []
                for cell in row:
                    if not cell.tag.endswith("c"):
                        continue
                    cell_type = cell.attrib.get("t")
                    value = ""
                    for child in cell:
                        if child.tag.endswith("v") and child.text:
                            value = child.text
                            break
                        if child.tag.endswith("is"):
                            value = "".join(child.itertext()).strip()
                            break
                    if cell_type == "s" and value.isdigit() and int(value) < len(shared):
                        value = shared[int(value)]
                    if value:
                        cells.append(value)
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _read_pptx(path: str) -> str:
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            name
            for name in zf.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
    return _zip_xml_text(path, members=slide_names)


def _read_open_document(path: str) -> str:
    return _zip_xml_text(path, members=["content.xml"])


def _read_epub(path: str) -> str:
    return _zip_xml_text(path, suffixes=(".xhtml", ".html", ".htm", ".xml"))


def _read_xls(path: str) -> str:
    """解析老版 Excel .xls。依赖本机安装 Microsoft Excel。"""
    try:
        import pythoncom
        import win32com.client as win32
    except Exception as exc:
        raise ValueError(".xls 解析需要 Windows + 已安装 Microsoft Excel，建议另存为 .xlsx。") from exc

    pythoncom.CoInitialize()
    excel = None
    try:
        excel = win32.gencache.EnsureDispatch("Excel.Application")
        excel.Visible = False
        wb = excel.Workbooks.Open(os.path.abspath(path), ReadOnly=True)
        parts: List[str] = []
        for sheet in wb.Worksheets:
            values = sheet.UsedRange.Value
            if values is None:
                continue
            if not isinstance(values, tuple):
                values = ((values,),)
            for row in values:
                if not isinstance(row, tuple):
                    row = (row,)
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        wb.Close(False)
        return "\n".join(parts)
    except Exception as exc:
        raise ValueError(f".xls 解析失败（请确认已安装 Microsoft Excel）：{exc}") from exc
    finally:
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _read_ppt(path: str) -> str:
    """解析老版 PowerPoint .ppt。依赖本机安装 Microsoft PowerPoint。"""
    try:
        import pythoncom
        import win32com.client as win32
    except Exception as exc:
        raise ValueError(".ppt 解析需要 Windows + 已安装 Microsoft PowerPoint，建议另存为 .pptx。") from exc

    pythoncom.CoInitialize()
    app = None
    try:
        app = win32.gencache.EnsureDispatch("PowerPoint.Application")
        presentation = app.Presentations.Open(os.path.abspath(path), WithWindow=False)
        parts: List[str] = []
        for slide in presentation.Slides:
            for shape in slide.Shapes:
                try:
                    if shape.HasTextFrame and shape.TextFrame.HasText:
                        text = shape.TextFrame.TextRange.Text
                        if text and text.strip():
                            parts.append(text.strip())
                except Exception:
                    continue
        presentation.Close()
        return "\n".join(parts)
    except Exception as exc:
        raise ValueError(f".ppt 解析失败（请确认已安装 Microsoft PowerPoint）：{exc}") from exc
    finally:
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _read_txt(path: str) -> str:
    # 优先 utf-8，失败回退 gbk（Windows 常见）
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法识别文件编码：{path}")


def load_file(
    path: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> List[Chunk]:
    """解析单个文件并切分为 Chunk 列表。"""
    chunk_size = chunk_size or settings.chunk_size
    chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap
    ext = os.path.splitext(path)[1].lower()
    filename = os.path.basename(path)

    if ext == ".pdf":
        raw = _read_pdf(path)
    elif ext == ".docx":
        raw = _read_docx(path)
    elif ext == ".doc":
        raw = _read_doc(path)
    elif ext == ".xlsx":
        raw = _read_xlsx(path)
    elif ext == ".xls":
        raw = _read_xls(path)
    elif ext == ".pptx":
        raw = _read_pptx(path)
    elif ext == ".ppt":
        raw = _read_ppt(path)
    elif ext in (".odt", ".ods", ".odp"):
        raw = _read_open_document(path)
    elif ext == ".epub":
        raw = _read_epub(path)
    elif ext == ".csv":
        raw = _read_csv(path)
    elif ext in (".json", ".jsonl"):
        raw = _read_json(path)
    elif ext == ".xml":
        raw = _read_xml(path)
    elif ext in (".html", ".htm"):
        raw = _read_html(path)
    elif ext == ".rtf":
        raw = _read_rtf(path)
    elif ext in (".txt", ".md"):
        raw = _read_txt(path)
    elif ext in SUPPORTED_TYPES["image_ocr"]:
        raw = _read_image_ocr(path)
    else:
        supported = sorted(e for exts in SUPPORTED_TYPES.values() for e in exts)
        raise ValueError(f"不支持的文件类型：{ext}（当前支持：{', '.join(supported)}）")

    text = clean_text(raw)
    if not text:
        logger.warning("文件无有效文本内容：%s", filename)
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=_SEPARATORS,
        keep_separator=True,
    )
    pieces = splitter.split_text(text)

    chunks = [
        Chunk(text=piece.strip(), source=filename, chunk_index=i)
        for i, piece in enumerate(pieces)
        if piece.strip()
    ]
    logger.info("文件 %s 解析完成，切分 %d 个片段", filename, len(chunks))
    return chunks


def load_files(
    paths: List[str],
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> LoadResult:
    """批量加载多个文件，并跨文件做内容去重。"""
    result = LoadResult()
    seen_hashes: set[str] = set()

    for path in paths:
        try:
            for chunk in load_file(path, chunk_size, chunk_overlap):
                if chunk.content_hash in seen_hashes:
                    result.skipped_duplicates += 1
                    continue
                seen_hashes.add(chunk.content_hash)
                result.chunks.append(chunk)
        except Exception as exc:
            logger.error("加载文件失败 %s: %s", path, exc)
            raise

    logger.info(
        "批量加载完成：有效片段 %d，去重跳过 %d",
        len(result.chunks),
        result.skipped_duplicates,
    )
    return result
